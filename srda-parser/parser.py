import re
from datetime import datetime

import pandas as pd
from docx import Document

from .models import ParsedTable, Data, Variable


def extract_title_from_document(doc: Document):
    # greedily extract the first paragraph
    for paragraph in doc.paragraphs:
        if len(paragraph.text) > 0:
            title = paragraph.text
            return title


def extract_date_from_document(doc: Document):
    # assume title is the first paragraph
    title = extract_title_from_document(doc)

    # extract date from title
    year = int(re.findall(r"(\d{2,3})年", title)[0]) + 1911
    return year


def extract_tables_from_document(doc: Document):
    tables = doc.tables
    return tables


def parse_docx_table(table):
    table_data = []

    # construct header from the first row
    header_cells = table.rows[0].cells
    header = []
    for cellidx in range(len(header_cells)):
        if cellidx > 0 and header_cells[cellidx].text == header_cells[cellidx - 1].text:
            pass  # left merge cells with duplicated header
        else:
            text = header_cells[cellidx].text.replace("\n", "")
            header.append(text)

    # init a variable to store running common title
    common_title = (None, None)  # (q_id, common_title)

    # loop through rows
    for row in table.rows[1:]:
        row_data = []

        # loop through cells within row
        for cell in row.cells:
            if len(row_data) > 0 and cell.text == row_data[-1]:
                pass  # left merge cells with duplicated content
            else:
                row_data.append(cell.text.strip())

        # truncate row_data to match header length
        row_data = row_data[: len(header)]

        # skip empty rows
        if len(row_data) == 0: continue

        # identify rows that served as common title for following rows
        # and rollover the title to the next row
        if row_data[0] != "" and len(row_data) > 1 and len(row_data) < len(header) - 1:
            # is a common title row
            common_title = (row_data[0], row_data[1])  # (q_id, common_title)
        else:
            # is not a common title row
            # get the position of variable description
            description_idx = header.index("變項說明")
            if row_data[0] == common_title[0]:
                # same q_id as common title
                row_data[description_idx] = common_title[1] + row_data[description_idx]

            # skip rows that served as separators
            if len(row_data) > 1:
                # right pad row_data to match header length
                row_data += [""] * (len(header) - len(row_data))

                table_data.append(row_data)

    pt = ParsedTable(header, table_data)
    return pt


def read_srda_docx(file_path, title=None, description=None, date=None):
    doc = Document(file_path)

    # construct Data object
    data = Data()
    data.title = title or extract_title_from_document(doc)
    data.description = description
    data.date = date or extract_date_from_document(doc)

    # extract variable informations from parsed codebook
    tables = extract_tables_from_document(doc)
    parsed_codebook = parse_docx_table(tables[0])  # first table is the codebook

    # construct Variable objects
    data.variables = parsed_codebook.to_variable_list()

    return data
